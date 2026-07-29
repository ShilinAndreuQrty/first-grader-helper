from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from app.knowledge.importer import normalize_text, parse_docx


def build_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("1. Адаптация")
    document.add_paragraph("Кто такой тьютор?", style="List Paragraph")
    document.add_paragraph("Тьютор помогает первокурснику.", style="List Paragraph")
    document.add_paragraph("Как найти расписание?", style="List Paragraph")
    document.add_paragraph("Откройте сайт.", style="List Paragraph")
    document.add_paragraph("Сохраните номер группы.", style="List Paragraph")
    document.add_paragraph("2. Учёба")
    document.add_paragraph(
        "Что такое зачёт?\nЗачёт — форма промежуточной аттестации.",
        style="List Paragraph",
    )
    document.save(path)


def test_importer_keeps_multiline_and_inline_answers(tmp_path: Path) -> None:
    source = tmp_path / "fixture.docx"
    build_fixture(source)

    result = parse_docx(source)

    assert result.report.categories == 2
    assert result.report.questions == 3
    assert "Сохраните номер группы." in result.entries[1].answer_markdown
    assert result.entries[2].answer_markdown == "Зачёт — форма промежуточной аттестации."
    assert result.entries[0].status == "published"
    assert normalize_text("Где расписание ?") == "где расписание?"


def test_importer_is_deterministic(tmp_path: Path) -> None:
    source = tmp_path / "fixture.docx"
    build_fixture(source)

    first = parse_docx(source)
    second = parse_docx(source)

    assert [entry.source_key for entry in first.entries] == [
        entry.source_key for entry in second.entries
    ]


def test_committed_seed_matches_source_baseline() -> None:
    seed_path = Path(__file__).parents[1] / "app" / "knowledge" / "seed" / "faq.json"
    payload = json.loads(seed_path.read_text(encoding="utf-8"))

    assert len(payload["categories"]) == 8
    assert len(payload["entries"]) == 60
    assert len(payload["report"]["duplicates"]) == 1
    assert sum(item["status"] == "published" for item in payload["entries"]) == 23
    assert sum(item["status"] == "archived" for item in payload["entries"]) == 1


def test_russian_normalization_handles_typos_preprocessing() -> None:
    assert normalize_text("  ТЬЮТОР — это...  ") == "тьютор - это"
    assert normalize_text("Ещё") == "еще"
